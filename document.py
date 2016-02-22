#!usr/bin/env python

print "importing libraries..."

# saving and displaying output
import logging
logging.basicConfig(format='%(asctime)s : %(levelname)s : %(message)s', level=logging.INFO)

# text conversion 
from subprocess import Popen, PIPE
from docx       import opendocx, getdocumenttext
#http://stackoverflow.com/questions/5725278/python-help-using-pdfminer-as-a-library
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout    import LAParams
from pdfminer.pdfpage   import PDFPage
from cStringIO          import StringIO

# proofreading by After The Deadline (https://blog.afterthedeadline.com/2009/09/15/python-bindings-for-atd/)
import ATD

# natural language processing
from nltk        import data, tokenize, pos_tag
from gensim      import corpora, models, similarities
from collections import defaultdict

# LDA visualization	
import pyLDAvis
import pyLDAvis.gensim

# database 
import json
from firebase import firebase

class Document():
	'''
	convert, parse, and operate on input text

	TODO: 
	- [ ] change all print statements to logs
	- [x] support for .pdf, .doc, .docx, and .odt
	- [x] support for .txt
	- [x] sentence tokenize
	- [ ] paragraph tokenize !!!
	- [x] part of speech tagging
	- [/] writing document to database -- what to include? everything?
	- [ ] LSA? lsa.colorado.edu?
	- [ ] word similarity matrix
	- [x] firebase
	- [ ] word2vec!!   https://radimrehurek.com/gensim/models/word2vec.html --> class gensim.models.word2vec.Word2Vec: Class for training, using and evaluating neural networks described in https://code.google.com/p/word2vec/
	- [ ] wikipedia dumps? (Already have AA through BC... might want to find precompiled set? takes many days to parse all pages from compressed archive)
	- [ ] bigram transformer (process phrases like words) !!
	- [ ] 
		.
		.
		.
	'''
	# establish connection to firebase
	database = firebase.FirebaseApplication('https://grademebaby.firebaseio.com/', None)

	# we might want to train the tokenizer on the format of the input text using PunktSentenceTokenizer(text)
	# these split the document into sentences
	tokenizer     = tokenize.TreebankWordTokenizer()
	sent_detector = data.load('tokenizers/punkt/english.pickle')

	def __init__(self, filename):
		self.filename     = filename # contains extension
		self.raw          = str()    # document as str 
		self.preprocessed = dict()   # text converted into ( word, lemma, [POS] ) format
		self.stats        = dict()

	def convert_pdf_to_txt(self, path):
		rsrcmgr     = PDFResourceManager()
		retstr      = StringIO()
		codec       = 'utf-8'
		laparams    = LAParams()
		device      = TextConverter(rsrcmgr, retstr, codec=codec, laparams=laparams)
		fp          = file(path, 'rb')
		interpreter = PDFPageInterpreter(rsrcmgr, device)
		password    = ""
		maxpages    = 0
		caching     = True
		pagenos     = set()

		map(interpreter.process_page, [page for page in PDFPage.get_pages(fp, pagenos, maxpages=maxpages, password=password,caching=caching, check_extractable=True)])
  
  		fp.close()
    		device.close()
    		str = retstr.getvalue()
    		retstr.close()
    		return str

	def document_to_text(self, filename, file_path):
		if filename[-4:] == ".doc":
        		cmd            = ['antiword', file_path]
        		p              = Popen(cmd, stdout=PIPE)
        		stdout, stderr = p.communicate()
        		self.raw       = stdout.decode('ascii', 'ignore')
    		
		elif filename[-5:] == ".docx":
        		document        = opendocx(file_path)
        		paratextlist    = getdocumenttext(document)
        		newparatextlist = []
        		for paratext in paratextlist:
        			 newparatextlist.append(paratext.encode("utf-8"))
       			self.raw = '\n\n'.join(newparatextlist)
    		
		elif filename[-4:] == ".odt":
        		cmd            = ['odt2txt', file_path]
        		p              = Popen(cmd, stdout=PIPE)
        		stdout, stderr = p.communicate()
        		self.raw       = stdout.decode('ascii', 'ignore')
    	
		elif filename[-4:] == ".pdf":
        		self.raw = self.convert_pdf_to_txt(file_path)
		
		elif filename[-4:] == ".txt":
			with open(file_path, 'r') as file_:
				self.raw = file_.read()

	def proofread(self):
		 # our API key for AfterTheDeadline
		ATD.setDefaultKey(hash("DoubleCheck")) 

		# check the document for grammar and spelling errors 	
		errors = ATD.checkDocument(self.raw)

		'''
		# print the errors
		for error in errors: 	
			print "%s error for: %s **%s**" % (error.type, error.precontext, error.string)
			print "some suggestions: %s" % (", ".join(error.suggestions),)
		'''

		# write the errors to the database
		err2db = [{"type":        error.type,
			   "precontext":  error.precontext,
			   "string":      error.string,
			   "suggestions": error.suggestions} for error in errors] 

		print json.dumps(err2db, sort_keys=True, indent=4)

	def vectorize(self):
		# tokenize and remove stopwords
		sentences = self.sent_detector.tokenize(self.raw.decode('utf-8').strip())
		stoplist  = set('for a of the and to in'.split())
		texts     = [[word for word in sentence.lower().split() if word not in stoplist] for sentence in sentences]
		
		# compute the frequency of each token
		frequency = defaultdict(int)
		for text in texts:
			for token in text:
				frequency[token] += 1

		# remove words that appear only once
		texts = [[token for token in text if frequency[token] > 1] for text in texts]
		
		# construct a gensim dictionary and corpus (bag of words)
		dictionary = corpora.Dictionary(texts)
		corpus     = [dictionary.doc2bow(text) for text in texts] # currently, "text" is a sentence in the document

		# define LDA model
		lda = models.ldamodel.LdaModel( corpus       = corpus, 
						id2word      = dictionary,
						num_topics   = 100, #what should this be ???
						update_every = 1, 
						chunksize    = 10000, 
						passes       = 100 )
		
		# print the extracted topics
		lda.print_topics(10)	

		# visualize the lda space
		vis_data = pyLDAvis.gensim.prepare(lda, corpus, dictionary)
        	pyLDAvis.display(vis_data)
       		pyLDAvis.show(vis_data)

	def preprocess_text(self):
		sentences = self.sent_detector.tokenize(self.raw.decode('utf-8').strip())
		tokens    = [self.tokenizer.tokenize(sentence) for sentence in sentences]		
		pos       = [pos_tag(token) for token in tokens]
		
		# final format includes 1) token, 2) lemma, and 3) list of part of speech tags		
		pos = [[(word, word, [postag]) for (word, postag) in sentence] for sentence in pos]		

		self.preprocessed = { 'sentences': sentences,
				      'tokens'   : tokens,
				      'pos'      : pos }

	def statistics(self):
		self.stats['sentences'] = len(self.preprocessed['sentences'])
		self.stats['tokens']    = len(self.preprocessed['tokens'])

def main():
	user     = raw_input('user: ')
	filename = raw_input('filename: ') 
	name     = filename[:-5] # this is wrong !!

	doc = Document(filename)

	# check to see if the file is already in the database
	user_files = doc.database.get('/documents/' + user, None)
	if name in user_files.keys():
		''' <-- currently broken?
		print "loading document from database..."
		doc.raw          = user_files[name].keys()[0]['raw'] # complicated syntax to bypass random key generation :(
		doc.preprocessed = user_files[name].keys()[0]['tagged'] 
		doc.stats        = user_files[name].keys()[0]['stats']
		'''
	else:
		print "converting document to raw text..."
		doc.document_to_text(doc.filename, doc.filename)
		print "proofreading the document..."
		doc.proofread()
		print "vectorizing raw text and performing LDA..."
		doc.vectorize() # must be called after document_to_test
		print "NOT preprocessing raw text..."
		#doc.preprocess_text()
		print "NOT getting document statistics..."
		#doc.statistics()
		print "NOT writing document to database..."
		'''
		db_entry = { "filename": doc.filename,
		     	     "raw":      doc.raw,
		     	     "tagged":   doc.preprocessed,
		     	     "stats":    doc.stats }	
		'''
		#doc.database.post('/documents/' + user + "/" + name, db_entry) 	

if __name__=="__main__": main()
